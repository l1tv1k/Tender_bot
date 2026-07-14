import re
import logging
from docx import Document
from typing import List, Optional

logger = logging.getLogger(__name__)


class SmartDocxExtractor:
    def __init__(self):
        # Триггеры для начала сбора текста (мясо)
        self.start_markers = [
            r"техническое\s+задание",
            r"описание\s+объекта\s+закупки",
            r"спецификация",
            r"требования\s+к\s+оказанию\s+услуг",
            r"информационная\s+карта"
        ]

        # Триггеры для остановки сбора (юридическая вода)
        self.stop_markers = [
            r"проект\s+контракта",
            r"проект\s+договора",
            r"обоснование\s+начальной\s+\(максимальной\)\s+цены",
            r"инструкция\s+по\s+заполнению",
            r"требования\s+к\s+содержанию\s+и\s+составу\s+заявки"
        ]

    def _is_marker(self, text: str, markers: List[str]) -> bool:
        """Проверяет, является ли строка заголовком-триггером."""
        clean_text = text.lower().strip()

        # Заголовок обычно короткий. Если абзац больше 200 символов — это просто текст,
        # даже если там встречается фраза "проект контракта".
        if len(clean_text) > 200:
            return False

        return any(re.search(marker, clean_text) for marker in markers)

    def _table_to_markdown(self, table) -> str:
        """
        Превращает таблицу из Word в Markdown-формат.
        Нейросети (особенно Mistral) великолепно понимают Markdown.
        Именно в таблицах лежат графики постов охраны и адреса объектов.
        """
        markdown_table = []
        for row in table.rows:
            # Читаем ячейки, убираем переносы строк внутри ячейки
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]

            # В Word часто бывают объединенные ячейки, которые дублируют текст. Очищаем:
            row_data = list(dict.fromkeys(row_data))

            # Пропускаем пустые строки
            if any(row_data):
                markdown_table.append(" | ".join(row_data))

        return "\n".join(markdown_table)

    def extract(self, file_path: str) -> str:
        """
        Главный метод извлечения. Итерируется по документу и забирает только суть.
        """
        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"Не удалось открыть {file_path}: {e}")
            return ""

        extracted_lines = []
        recording = False
        target_found = False

        # Читаем все абзацы документа
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Проверяем, не начался ли нужный нам раздел
            if not recording and self._is_marker(text, self.start_markers):
                recording = True
                target_found = True
                extracted_lines.append(f"\n### {text.upper()} ###\n")
                continue

            # Проверяем, не пора ли остановиться
            if recording and self._is_marker(text, self.stop_markers):
                recording = False
                continue

            # Если мы в нужной зоне, записываем текст
            if recording:
                extracted_lines.append(text)

        # Если не нашли явных заголовков ТЗ (бывают кривые документы),
        # берем весь текст как Fallback (резервный план)
        if not target_found:
            logger.warning(f"Триггеры не найдены в {file_path}. Извлекаем весь текст.")
            extracted_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Обязательно добавляем таблицы (в них самая ценная информация)
        # Для простоты добавляем их в конец или после текста
        if doc.tables:
            extracted_lines.append("\n### ТАБЛИЦЫ ДОКУМЕНТА (ГРАФИКИ, АДРЕСА, ЦЕНЫ) ###\n")
            for table in doc.tables:
                extracted_lines.append(self._table_to_markdown(table))
                extracted_lines.append("\n---\n")

        # Собираем всё в одну строку
        final_text = "\n".join(extracted_lines)

        # Финальная очистка от лишних пустых строк
        final_text = re.sub(r'\n{3,}', '\n\n', final_text)

        return final_text